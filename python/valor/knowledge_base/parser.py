"""Document parser: PDF -> ParsedDocument (Word/Excel/TXT/MD in Task 2.2).

License: GPL-3.0-or-later WITH GPL-3.0-NonCommercial
"""
from __future__ import annotations

from dataclasses import dataclass, field
from pathlib import Path

import re as _re

import pdfplumber


@dataclass
class ParsedPage:
    page_no: int
    text: str


@dataclass
class ParsedTable:
    page_no: int
    rows: list[list[str]]
    caption: str | None = None


@dataclass
class HeadingNode:
    level: int
    text: str
    page_no: int
    children: list["HeadingNode"] = field(default_factory=list)


@dataclass
class ParsedDocument:
    file_path: str
    mime_type: str
    pages: list[ParsedPage] = field(default_factory=list)
    tables: list[ParsedTable] = field(default_factory=list)
    heading_tree: list[HeadingNode] = field(default_factory=list)
    full_text: str = ""


def parse(file_path: Path, mime_type: str) -> ParsedDocument:
    """Dispatch parser by mime_type."""
    if mime_type == "application/pdf":
        return parse_pdf(file_path)
    if mime_type in (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/msword",
    ):
        return parse_word(file_path)
    if mime_type in (
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        "application/vnd.ms-excel",
    ):
        return parse_excel(file_path)
    if mime_type in ("text/plain", "text/markdown", "text/csv"):
        return parse_text(file_path, mime_type)
    raise ValueError(f"unsupported mime_type: {mime_type}")


def parse_pdf(file_path: Path) -> ParsedDocument:
    """Parse PDF with pdfplumber: extract text + tables per page."""
    doc = ParsedDocument(file_path=str(file_path), mime_type="application/pdf")
    text_parts: list[str] = []
    # Track last heading text across pages for table caption inference
    last_heading: str | None = None
    with pdfplumber.open(file_path) as pdf:
        for idx, page in enumerate(pdf.pages, start=1):
            page_text = page.extract_text() or ""
            doc.pages.append(ParsedPage(page_no=idx, text=page_text))
            text_parts.append(page_text)
            # Extract tables with caption inferred from nearest preceding heading
            for tbl in page.extract_tables() or []:
                doc.tables.append(ParsedTable(page_no=idx, rows=tbl, caption=last_heading))
            # Heading heuristic: chars with size >= 14 and line length <= 30
            new_headings = _extract_headings_from_page_returned(page, idx, doc.heading_tree)
            if new_headings:
                last_heading = new_headings[-1].text
    doc.full_text = "\n\n".join(text_parts)
    return doc


def _extract_headings_from_page(page, page_no: int, tree: list[HeadingNode]) -> None:
    """Heuristic: lines with font size >= 14 and char count <= 30 are headings."""
    _extract_headings_from_page_returned(page, page_no, tree)


def _extract_headings_from_page_returned(page, page_no: int, tree: list[HeadingNode]) -> list[HeadingNode]:
    """Same as _extract_headings_from_page but returns the new headings (for caption tracking)."""
    new_headings: list[HeadingNode] = []
    try:
        words = page.extract_words(extra_attrs=["size"])
    except Exception:
        return new_headings
    # Group words by line (y0 rounded)
    lines: dict[float, list[dict]] = {}
    for w in words:
        key = round(w.get("top", 0), 0)
        lines.setdefault(key, []).append(w)
    for y in sorted(lines.keys()):
        words_in_line = sorted(lines[y], key=lambda w: w.get("x0", 0))
        text = "".join(w.get("text", "") for w in words_in_line).strip()
        if not text or len(text) > 30:
            continue
        avg_size = sum(w.get("size", 0) for w in words_in_line) / len(words_in_line)
        if avg_size >= 14:
            node = HeadingNode(level=1, text=text, page_no=page_no)
            tree.append(node)
            new_headings.append(node)
    return new_headings


def parse_word(file_path: Path) -> ParsedDocument:
    """Parse Word .docx with python-docx."""
    from docx import Document as DocxDocument

    doc = ParsedDocument(
        file_path=str(file_path),
        mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    d = DocxDocument(str(file_path))
    text_parts: list[str] = []
    for para in d.paragraphs:
        text = para.text.strip()
        if text:
            text_parts.append(text)
            # Heading detection by style
            if para.style and para.style.name and para.style.name.startswith("Heading"):
                try:
                    level = int(para.style.name.split()[-1])
                except ValueError:
                    level = 1
                doc.heading_tree.append(HeadingNode(level=level, text=text, page_no=1))
    # Tables
    for tbl in d.tables:
        rows = [[cell.text for cell in row.cells] for row in tbl.rows]
        doc.tables.append(ParsedTable(page_no=1, rows=rows))
    doc.full_text = "\n\n".join(text_parts)
    doc.pages.append(ParsedPage(page_no=1, text=doc.full_text))
    return doc


def parse_excel(file_path: Path) -> ParsedDocument:
    """Parse Excel .xlsx with openpyxl: each sheet -> markdown table."""
    import openpyxl

    doc = ParsedDocument(
        file_path=str(file_path),
        mime_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )
    wb = openpyxl.load_workbook(str(file_path), read_only=True, data_only=True)
    text_parts: list[str] = []
    for ws in wb.worksheets:
        text_parts.append(f"## {ws.title}\n")
        rows: list[list[str]] = []
        for row in ws.iter_rows(values_only=True):
            cells = [str(c) if c is not None else "" for c in row]
            if any(cells):
                rows.append(cells)
                text_parts.append("| " + " | ".join(cells) + " |")
        if rows:
            doc.tables.append(ParsedTable(page_no=1, rows=rows, caption=ws.title))
    wb.close()
    doc.full_text = "\n".join(text_parts)
    doc.pages.append(ParsedPage(page_no=1, text=doc.full_text))
    return doc


def parse_text(file_path: Path, mime_type: str) -> ParsedDocument:
    """Parse plain text / markdown."""
    doc = ParsedDocument(file_path=str(file_path), mime_type=mime_type)
    text = file_path.read_text(encoding="utf-8")
    doc.full_text = text
    doc.pages.append(ParsedPage(page_no=1, text=text))
    if mime_type == "text/markdown":
        for line in text.splitlines():
            if line.startswith("#"):
                level = len(line) - len(line.lstrip("#"))
                doc.heading_tree.append(
                    HeadingNode(level=level, text=line.lstrip("# ").strip(), page_no=1)
                )
    return doc


# ---------------------------------------------------------------------------
# Metadata extraction (Task 2.3)
# ---------------------------------------------------------------------------

_CN_NUM_MAP = {"一": 1, "二": 2, "三": 3, "四": 4, "1": 1, "2": 2, "3": 3, "4": 4}

_DATE_PATTERNS = [
    _re.compile(r"(\d{4})\s*年\s*(\d{1,2})\s*月\s*(\d{1,2})\s*日"),
    _re.compile(r"(\d{4})-(\d{2})-(\d{2})"),
    _re.compile(r"(\d{4})/(\d{1,2})/(\d{1,2})"),
]

_TICKER_PATTERN = _re.compile(r"(?:股票代码|代码|证券代码)?[：:\s]*(6\d{5}|0\d{5}|3\d{5}|8\d{5}|4\d{5})")

_PERIOD_PATTERNS = [
    (_re.compile(r"(\d{4})\s*年第([一二三四1234])季度"), lambda m: f"{m.group(1)}Q{_CN_NUM_MAP[m.group(2)]}"),
    (_re.compile(r"(\d{4})Q([1-4])"), lambda m: f"{m.group(1)}Q{m.group(2)}"),
    (_re.compile(r"(\d{4})-(\d{2})-(\d{2})"), lambda m: f"{m.group(1)}-{m.group(2)}-{m.group(3)}"),
]


def extract_publish_date(parsed: ParsedDocument) -> str | None:
    """Extract publish date from first page text. Returns YYYY-MM-DD or None."""
    if not parsed.pages:
        return None
    text = parsed.pages[0].text
    for pat in _DATE_PATTERNS:
        m = pat.search(text)
        if m:
            y, mo, d = m.groups()
            return f"{int(y):04d}-{int(mo):02d}-{int(d):02d}"
    return None


def extract_ticker(parsed: ParsedDocument) -> str | None:
    """Extract A-share ticker (6-digit code) from first page."""
    if not parsed.pages:
        return None
    m = _TICKER_PATTERN.search(parsed.pages[0].text)
    return m.group(1) if m else None


def extract_report_period(parsed: ParsedDocument) -> str | None:
    """Extract report period like 2024Q3 or 2024-09-30."""
    if not parsed.pages:
        return None
    text = parsed.pages[0].text
    for pat, fmt in _PERIOD_PATTERNS:
        m = pat.search(text)
        if m:
            return fmt(m)
    return None